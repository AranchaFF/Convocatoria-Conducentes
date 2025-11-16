import streamlit as st
import pandas as pd
import io
import base64
from datetime import datetime, date
import openpyxl
from PIL import Image
import pytesseract
import PyPDF2
import re
import platform
import os
import docx

# Configuración de Tesseract OCR
if platform.system() == 'Windows':
    posibles_rutas = [
        r'C:\Program Files\Tesseract-OCR\tesseract.exe',
        r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
        r'C:\Tesseract-OCR\tesseract.exe'
    ]
    for ruta in posibles_rutas:
        if os.path.exists(ruta):
            pytesseract.pytesseract.tesseract_cmd = ruta
            break

st.set_page_config(
    page_title="Smartmind - Documentación Convocatoria",
    page_icon="",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Diseño Moderno Profesional
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600;700&display=swap');
    
    :root {
        /* Paleta Moderna */
        --primary: #2563eb;
        --primary-light: #3b82f6;
        --primary-dark: #1d4ed8;
        --accent: #8b5cf6;
        --success: #10b981;
        --warning: #f59e0b;
        --error: #ef4444;
        
        /* Neutrales */
        --gray-50: #f9fafb;
        --gray-100: #f3f4f6;
        --gray-200: #e5e7eb;
        --gray-300: #d1d5db;
        --gray-400: #9ca3af;
        --gray-500: #6b7280;
        --gray-600: #4b5563;
        --gray-700: #374151;
        --gray-800: #1f2937;
        --gray-900: #111827;
        
        /* Superficies */
        --bg-primary: #ffffff;
        --bg-secondary: var(--gray-50);
        --border-color: var(--gray-200);
        
        /* Sombras */
        --shadow-sm: 0 1px 2px 0 rgb(0 0 0 / 0.05);
        --shadow-md: 0 4px 6px -1px rgb(0 0 0 / 0.1);
        --shadow-lg: 0 10px 15px -3px rgb(0 0 0 / 0.1);
        --shadow-xl: 0 20px 25px -5px rgb(0 0 0 / 0.1);
    }
    
    /* Reset Global */
    * {
        font-family: 'Outfit', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
        transition: all 0.2s ease;
    }
    
    /* Ocultar elementos de Streamlit */
    #MainMenu, footer, header {visibility: hidden;}
    
    /* Fondo de la app */
    .stApp {
        background: linear-gradient(to bottom right, var(--gray-50), var(--gray-100));
    }
    
    /* Contenedor principal */
    .main .block-container {
        padding: 2rem;
        max-width: 1280px;
        margin: 0 auto;
    }
    
    /* === TIPOGRAFÍA === */
    h1 {
        font-size: 2rem !important;
        font-weight: 700 !important;
        color: var(--gray-900) !important;
        margin-bottom: 0.5rem !important;
        letter-spacing: -0.025em !important;
    }
    
    h2 {
        font-size: 1.5rem !important;
        font-weight: 600 !important;
        color: var(--gray-800) !important;
        margin: 2rem 0 1rem 0 !important;
        letter-spacing: -0.02em !important;
    }
    
    h3 {
        font-size: 1.125rem !important;
        font-weight: 600 !important;
        color: var(--gray-700) !important;
        margin-bottom: 0.75rem !important;
    }
    
    p {
        color: var(--gray-600);
        line-height: 1.6;
        font-size: 0.9375rem;
    }
    
    /* === HEADER MODERNO === */
    .header-container {
        background: linear-gradient(135deg, var(--primary) 0%, var(--primary-dark) 100%);
        padding: 2.5rem;
        border-radius: 16px;
        margin-bottom: 2rem;
        box-shadow: var(--shadow-xl);
        position: relative;
        overflow: hidden;
    }
    
    .header-container::before {
        content: '';
        position: absolute;
        top: -50%;
        right: -20%;
        width: 500px;
        height: 500px;
        background: radial-gradient(circle, rgba(255,255,255,0.1) 0%, transparent 70%);
        border-radius: 50%;
    }
    
    .header-title {
        color: white;
        font-size: 2rem;
        font-weight: 700;
        margin: 0;
        position: relative;
        z-index: 1;
    }
    
    .header-subtitle {
        color: rgba(255,255,255,0.95);
        font-size: 1rem;
        margin-top: 0.5rem;
        font-weight: 400;
        position: relative;
        z-index: 1;
    }
    
    /* === CARDS MODERNOS === */
    .custom-card {
        background: var(--bg-primary);
        border: 1px solid var(--border-color);
        border-radius: 12px;
        padding: 1.5rem;
        margin-bottom: 1rem;
        box-shadow: var(--shadow-sm);
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
    }
    
    .custom-card:hover {
        box-shadow: var(--shadow-md);
        transform: translateY(-2px);
        border-color: var(--primary-light);
    }
    
    .custom-card h3 {
        color: var(--gray-900);
        font-size: 1rem;
        font-weight: 600;
        margin: 0 0 0.5rem 0;
    }
    
    .custom-card p {
        color: var(--gray-600);
        font-size: 0.875rem;
        margin: 0;
        line-height: 1.5;
    }
    
    /* === TABS MODERNOS === */
    .stTabs [data-baseweb="tab-list"] {
        gap: 0.5rem;
        border-bottom: 2px solid var(--border-color);
        background: transparent;
    }
    
    .stTabs [data-baseweb="tab"] {
        background: transparent;
        border: none;
        color: var(--gray-500);
        font-weight: 500;
        padding: 0.75rem 1.25rem;
        font-size: 0.9375rem;
        border-radius: 8px 8px 0 0;
    }
    
    .stTabs [data-baseweb="tab"]:hover {
        background: var(--gray-100);
        color: var(--gray-700);
    }
    
    .stTabs [aria-selected="true"] {
        color: var(--primary) !important;
        background: var(--bg-primary);
        border-bottom: 3px solid var(--primary);
        font-weight: 600;
    }
    
    .stTabs [data-baseweb="tab-panel"] {
        background: var(--bg-primary);
        border: 1px solid var(--border-color);
        border-top: none;
        border-radius: 0 0 12px 12px;
        padding: 2rem;
        box-shadow: var(--shadow-sm);
    }
    
    /* === FILE UPLOADER === */
    .stFileUploader {
        background: var(--bg-primary);
        border: 2px dashed var(--border-color);
        border-radius: 12px;
        padding: 2rem;
        transition: all 0.3s ease;
    }
    
    .stFileUploader:hover {
        border-color: var(--primary);
        background: var(--bg-secondary);
    }
    
    .stFileUploader label {
        color: var(--gray-700) !important;
        font-weight: 500 !important;
        font-size: 0.9375rem !important;
    }
    
    /* === BOTONES MODERNOS === */
    .stButton button {
        background: linear-gradient(135deg, var(--primary) 0%, var(--primary-light) 100%);
        color: white;
        font-weight: 600;
        border: none;
        border-radius: 10px;
        padding: 0.75rem 2rem;
        font-size: 0.9375rem;
        box-shadow: var(--shadow-md);
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
    }
    
    .stButton button:hover {
        transform: translateY(-2px);
        box-shadow: var(--shadow-lg);
    }
    
    .stButton button:active {
        transform: translateY(0);
    }
    
    .stDownloadButton button {
        background: linear-gradient(135deg, var(--success) 0%, #059669 100%);
        color: white;
        font-weight: 600;
        border: none;
        border-radius: 10px;
        padding: 0.75rem 2rem;
        font-size: 0.9375rem;
        box-shadow: var(--shadow-md);
    }
    
    .stDownloadButton button:hover {
        transform: translateY(-2px);
        box-shadow: var(--shadow-lg);
    }
    
    /* === MÉTRICAS === */
    [data-testid="stMetric"] {
        background: var(--bg-primary);
        border: 1px solid var(--border-color);
        border-radius: 12px;
        padding: 1.25rem;
        box-shadow: var(--shadow-sm);
        transition: all 0.3s ease;
    }
    
    [data-testid="stMetric"]:hover {
        box-shadow: var(--shadow-md);
        transform: translateY(-2px);
    }
    
    [data-testid="stMetricValue"] {
        font-size: 2rem !important;
        font-weight: 700 !important;
        color: var(--primary) !important;
        background: linear-gradient(135deg, var(--primary) 0%, var(--accent) 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
    }
    
    [data-testid="stMetricLabel"] {
        font-size: 0.75rem !important;
        font-weight: 600 !important;
        color: var(--gray-500) !important;
        text-transform: uppercase;
        letter-spacing: 0.05em;
    }
    
    /* === PROGRESS BAR === */
    .stProgress > div > div {
        background: var(--gray-200);
        border-radius: 10px;
        height: 8px;
    }
    
    .stProgress > div > div > div {
        background: linear-gradient(90deg, var(--primary) 0%, var(--accent) 100%);
        border-radius: 10px;
    }
    
    /* === ALERTAS === */
    .stAlert {
        border-radius: 12px;
        border: 1px solid;
        padding: 1rem 1.25rem;
        font-size: 0.9375rem;
        box-shadow: var(--shadow-sm);
    }
    
    .stSuccess {
        background: linear-gradient(135deg, #ecfdf5 0%, #d1fae5 100%);
        color: #065f46;
        border-color: #86efac;
    }
    
    .stWarning {
        background: linear-gradient(135deg, #fffbeb 0%, #fef3c7 100%);
        color: #92400e;
        border-color: #fcd34d;
    }
    
    .stError {
        background: linear-gradient(135deg, #fef2f2 0%, #fee2e2 100%);
        color: #991b1b;
        border-color: #fca5a5;
    }
    
    .stInfo {
        background: linear-gradient(135deg, #eff6ff 0%, #dbeafe 100%);
        color: #1e40af;
        border-color: #93c5fd;
    }
    
    /* === RESULTADO === */
    .result-container {
        background: linear-gradient(135deg, #ecfdf5 0%, #d1fae5 100%);
        border: 2px solid #86efac;
        border-radius: 16px;
        padding: 2rem;
        margin: 2rem 0;
        box-shadow: var(--shadow-lg);
        position: relative;
        overflow: hidden;
    }
    
    .result-container::before {
        content: '';
        position: absolute;
        top: -50%;
        right: -20%;
        width: 400px;
        height: 400px;
        background: radial-gradient(circle, rgba(16, 185, 129, 0.1) 0%, transparent 70%);
        border-radius: 50%;
    }
    
    .result-title {
        color: #065f46;
        font-size: 1.5rem;
        font-weight: 700;
        margin-bottom: 0.75rem;
        position: relative;
        z-index: 1;
    }
    
    /* === SIDEBAR === */
    [data-testid="stSidebar"] {
        background: var(--bg-primary);
        border-right: 1px solid var(--border-color);
        box-shadow: var(--shadow-md);
    }
    
    .sidebar-title {
        color: var(--gray-900);
        font-size: 1.125rem;
        font-weight: 600;
    }
    
    /* === INPUTS === */
    .stTextInput input,
    .stTextArea textarea,
    .stSelectbox select {
        border: 1px solid var(--border-color);
        border-radius: 8px;
        padding: 0.625rem 0.875rem;
        font-size: 0.9375rem;
        background: var(--bg-primary);
        transition: all 0.2s ease;
    }
    
    .stTextInput input:focus,
    .stTextArea textarea:focus,
    .stSelectbox select:focus {
        border-color: var(--primary);
        box-shadow: 0 0 0 3px rgba(37, 99, 235, 0.1);
        outline: none;
    }
    
    /* === DIVISOR === */
    hr {
        border: none;
        border-top: 1px solid var(--border-color);
        margin: 2rem 0;
    }
    
    /* === SCROLLBAR === */
    ::-webkit-scrollbar {
        width: 8px;
        height: 8px;
    }
    
    ::-webkit-scrollbar-track {
        background: var(--gray-100);
    }
    
    ::-webkit-scrollbar-thumb {
        background: var(--gray-300);
        border-radius: 4px;
    }
    
    ::-webkit-scrollbar-thumb:hover {
        background: var(--gray-400);
    }
</style>
""", unsafe_allow_html=True)


# Funciones de extracción de texto de documentos
def extraer_texto_pdf(file):
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        texto = ""
        for page in pdf_reader.pages:
            texto += page.extract_text() + "\n"
        return texto
    except Exception as e:
        st.error(f"Error al leer PDF: {str(e)}")
        return ""

def extraer_texto_imagen(file):
    try:
        image = Image.open(file)
        texto = pytesseract.image_to_string(image, lang='spa')
        return texto
    except Exception as e:
        st.error(f"Error al procesar imagen: {str(e)}")
        return ""

def extraer_texto_word(file):
    try:
        doc = docx.Document(file)
        texto = ""
        for paragraph in doc.paragraphs:
            texto += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texto += cell.text + " "
            texto += "\n"
        return texto
    except Exception as e:
        st.error(f"Error al leer Word: {str(e)}")
        return ""

def extraer_texto_excel(file):
    """Extrae texto de un archivo Excel"""
    try:
        xls = pd.ExcelFile(file)
        texto = ""
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(file, sheet_name=sheet_name)
            texto += f"\n--- Hoja: {sheet_name} ---\n"
            texto += df.to_string(index=False) + "\n"
        return texto
    except Exception as e:
        st.error(f"Error al leer Excel: {str(e)}")
        return ""

def procesar_documento(file):
    """Procesa cualquier tipo de documento"""
    if file.type == "application/pdf":
        return extraer_texto_pdf(file)
    elif file.type in ["image/png", "image/jpeg", "image/jpg", "image/bmp", "image/tiff", "image/gif"]:
        return extraer_texto_imagen(file)
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extraer_texto_word(file)
    elif file.type in ["application/vnd.ms-excel", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"]:
        return extraer_texto_excel(file)
    else:
        st.warning(f"Tipo de archivo: {file.type}")
        try:
            return file.read().decode('utf-8')
        except:
            return ""

# NUEVA FUNCIÓN: Leer datos del Excel CTRL
def leer_datos_ctrl(excel_file):
    """Lee la pestaña CTRL del Excel CTRL de Alumnos"""
    datos_ctrl = {}

    try:
        excel_file.seek(0)

        if "CTRL" not in pd.ExcelFile(excel_file).sheet_names:
            st.warning("No se encontró la pestaña 'CTRL' en el Excel CTRL")
            return datos_ctrl

        st.write("Leyendo pestaña CTRL del Excel CTRL de Alumnos...")

        df_ctrl = pd.read_excel(excel_file, sheet_name="CTRL")
        st.write(f"Columnas encontradas en CTRL: {list(df_ctrl.columns)}")

        # Identificar columnas
        col_nombre = None
        col_dni = None
        col_corporacion = None
        col_baja = None
        col_motivo = None
        col_baja_ocupacion = None
        col_fecha_incorporacion = None

        for col in df_ctrl.columns:
            col_lower = str(col).lower().strip()

            st.write(f"   Analizando columna: '{col}' (lower: '{col_lower}')")

            if 'nombre' in col_lower or 'alumno' in col_lower:
                col_nombre = col
                st.write(f"     Identificada como NOMBRE")
            elif 'dni' in col_lower or 'nif' in col_lower:
                col_dni = col
                st.write(f"     Identificada como DNI")
            # VERIFICAR FECHA DE INCORPORACIÓN ANTES QUE CORPORACIÓN
            elif ('incorporacion' in col_lower or 'incorporación' in col_lower) and 'fecha' in col_lower:
                col_fecha_incorporacion = col
                st.write(f"     Identificada como FECHA INCORPORACIÓN")
            elif 'corporacion' in col_lower or 'corporación' in col_lower:
                col_corporacion = col
                st.write(f"     Identificada como CORPORACIÓN")
            elif 'motivo' in col_lower:
                col_motivo = col
                st.write(f"     Identificada como MOTIVO")
            elif ('baja' in col_lower and ('ocupacion' in col_lower or 'ocupación' in col_lower)) or \
                 ('ocupacion' in col_lower or 'ocupación' in col_lower) and '%' in col_lower:
                col_baja_ocupacion = col
                st.write(f"     Identificada como BAJA OCUPACIÓN")
            # PRIORIZAR "BAJA (FECHA)" sobre otras columnas con BAJA
            elif 'baja' in col_lower and 'fecha' in col_lower:
                col_baja = col
                st.write(f"     Identificada como BAJA (con fecha)")
            elif 'baja' in col_lower and col_baja is None:  # Solo si no se ha identificado una columna de baja aún
                st.write(f"     Columna con BAJA pero sin FECHA, ignorando: '{col}'")

        st.write(f"Columnas identificadas: Nombre={col_nombre}, DNI={col_dni}, Corporación={col_corporacion}, Baja={col_baja}, Motivo={col_motivo}, % Baja Ocupación={col_baja_ocupacion}, Fecha Incorporación={col_fecha_incorporacion}")

        # Procesar cada alumno
        for _, row in df_ctrl.iterrows():
            if col_nombre and pd.notna(row.get(col_nombre)):
                nombre = str(row[col_nombre]).strip().upper()

                if nombre:
                    # Procesar fecha - SOLO EXTRAER DIA, MES, AÑO
                    fecha_incorporacion_valor = None

                    st.write(f"   Procesando alumno: '{nombre}'")
                    st.write(f"    - col_fecha_incorporacion = {col_fecha_incorporacion}")

                    if col_fecha_incorporacion:
                        fecha_val = row.get(col_fecha_incorporacion)
                        st.write(f"    - Valor raw de fecha: {fecha_val}")
                        st.write(f"    - Tipo: {type(fecha_val)}")
                        st.write(f"    - pd.notna: {pd.notna(fecha_val)}")

                        if pd.notna(fecha_val):
                            if isinstance(fecha_val, pd.Timestamp):
                                fecha_incorporacion_valor = date(fecha_val.year, fecha_val.month, fecha_val.day)
                                st.write(f"     Convertido de Timestamp: {fecha_incorporacion_valor}")
                            elif isinstance(fecha_val, datetime):
                                fecha_incorporacion_valor = date(fecha_val.year, fecha_val.month, fecha_val.day)
                                st.write(f"     Convertido de datetime: {fecha_incorporacion_valor}")
                            else:
                                fecha_incorporacion_valor = str(fecha_val)
                                st.write(f"     Convertido a string: {fecha_incorporacion_valor}")
                        else:
                            st.write(f"     Fecha es NaN/None")
                    else:
                        st.write(f"     col_fecha_incorporacion es None")

                    # Procesar fecha de baja (convertir a date si es datetime)
                    baja_valor = None
                    if col_baja and pd.notna(row.get(col_baja)):
                        baja_raw = row[col_baja]
                        if isinstance(baja_raw, pd.Timestamp):
                            baja_valor = date(baja_raw.year, baja_raw.month, baja_raw.day)
                        elif isinstance(baja_raw, datetime):
                            baja_valor = date(baja_raw.year, baja_raw.month, baja_raw.day)
                        else:
                            baja_valor = str(baja_raw)

                    datos_ctrl[nombre] = {
                        "dni": str(row[col_dni]) if col_dni and pd.notna(row.get(col_dni)) else "",
                        "corporacion_a_clase": str(row[col_corporacion]) if col_corporacion and pd.notna(row.get(col_corporacion)) else "",
                        "baja": baja_valor if baja_valor else "",
                        "motivo": str(row[col_motivo]) if col_motivo and pd.notna(row.get(col_motivo)) else "",
                        "motivo_sin_parentesis": re.sub(r'\s*\([^)]*\)', '', str(row[col_motivo])).strip() if col_motivo and pd.notna(row.get(col_motivo)) else "",
                        "baja_ocupacion": str(row[col_baja_ocupacion]) if col_baja_ocupacion and pd.notna(row.get(col_baja_ocupacion)) else "",
                        "fecha_incorporacion": fecha_incorporacion_valor
                    }
                    st.write(f"   {nombre}: Fecha incorporación={fecha_incorporacion_valor}, Baja={baja_valor}")

        st.success(f" Datos del CTRL leídos: {len(datos_ctrl)} alumnos encontrados")
        return datos_ctrl

    except Exception as e:
        st.error(f"Error al leer Excel CTRL: {str(e)}")
        import traceback
        st.code(traceback.format_exc())
        return datos_ctrl

# Función para leer datos del Excel (pestañas CALIFICACIONES y ASISTENCIA)
def leer_datos_excel(excel_file, datos_evaluacion=None):
    """Lee las pestañas CALIFICACIONES y ASISTENCIA del Excel, y combina con evaluaciones"""
    datos = {
        "alumnos": {},
    }

    try:
        excel_file.seek(0)

        # ===== LEER CALIFICACIONES =====
        if "CALIFICACIONES" in pd.ExcelFile(excel_file).sheet_names:
            df_calif_raw = pd.read_excel(excel_file, sheet_name="CALIFICACIONES", header=None)
            st.write("Leyendo pestaña CALIFICACIONES...")

            fila_alumnos = None
            for idx, row in df_calif_raw.iterrows():
                if 'ALUMNADO' in str(row.values).upper() or 'ALUMNO' in str(row.values).upper():
                    fila_alumnos = idx
                    st.write(f"   Fila de datos de alumnos: {fila_alumnos}")
                    break

            if fila_alumnos:
                df_calif = pd.read_excel(excel_file, sheet_name="CALIFICACIONES", header=fila_alumnos)

                for _, row in df_calif.iterrows():
                    nombre_col = [col for col in df_calif.columns if 'ALUMN' in str(col).upper()]
                    if nombre_col and pd.notna(row.get(nombre_col[0])):
                        nombre = str(row[nombre_col[0]]).strip()
                        if nombre and 'ALUMN' not in nombre.upper():
                            dni_col = [col for col in df_calif.columns if 'DNI' in str(col).upper()]
                            dni = str(row[dni_col[0]]) if dni_col and pd.notna(row.get(dni_col[0])) else ""

                            datos["alumnos"][nombre] = {
                                "dni": dni,
                                "fcoo03": "",
                                "modulos_mf": {},
                                "porcentaje_asistencia": ""
                            }

        # ===== LEER ASISTENCIA =====
        excel_file.seek(0)
        if "ASISTENCIA" in pd.ExcelFile(excel_file).sheet_names:
            df_asist_raw = pd.read_excel(excel_file, sheet_name="ASISTENCIA", header=None)
            st.write("Leyendo pestaña ASISTENCIA...")

            fila_modulos = None
            for idx, row in df_asist_raw.iterrows():
                row_text = ' '.join([str(x) for x in row.values if pd.notna(x)])
                if 'MF0969' in row_text and 'MF0970' in row_text:
                    fila_modulos = idx
                    st.write(f"   Fila de módulos: {fila_modulos}")
                    break

            fila_alumnos = None
            for idx, row in df_asist_raw.iterrows():
                if 'ALUMNADO' in str(row.values).upper() or 'ALUMNO' in str(row.values).upper():
                    fila_alumnos = idx
                    st.write(f"   Fila de datos de alumnos: {fila_alumnos}")
                    break

            if fila_alumnos and fila_modulos:
                df_asist = pd.read_excel(excel_file, sheet_name="ASISTENCIA", header=fila_alumnos)
                fila_modulos_data = df_asist_raw.iloc[fila_modulos]

                st.write("**Mapeo de módulos encontrados:**")
                mapeo_modulos = {}
                for col_idx, valor in enumerate(fila_modulos_data):
                    if pd.notna(valor):
                        valor_str = str(valor)
                        if 'MF' in valor_str and '_' in valor_str:
                            match = re.search(r'(MF\d{4}_\d)', valor_str)
                            if match:
                                codigo_modulo = match.group(1)
                                if col_idx < len(df_asist.columns):
                                    nombre_col = df_asist.columns[col_idx]
                                    mapeo_modulos[nombre_col] = codigo_modulo
                                    st.write(f"   Columna '{nombre_col}' → {codigo_modulo}")

                        if 'FCOO' in valor_str:
                            if col_idx < len(df_asist.columns):
                                nombre_col = df_asist.columns[col_idx]
                                mapeo_modulos[nombre_col] = 'FCOO03'
                                st.write(f"   Columna '{nombre_col}' → FCOO03")

                for _, row in df_asist.iterrows():
                    nombre_col = [col for col in df_asist.columns if 'ALUMN' in str(col).upper()]
                    if nombre_col and pd.notna(row.get(nombre_col[0])):
                        nombre = str(row[nombre_col[0]]).strip()
                        if nombre in datos["alumnos"]:
                            # Buscar la ÚLTIMA columna que tenga % y un valor numérico válido
                            porcentaje_encontrado = False
                            columnas_con_porcentaje = []

                            for col in df_asist.columns:
                                col_str = str(col)
                                if '%' in col_str:
                                    valor = row.get(col)
                                    # Solo considerar si es un número válido entre 0 y 1 (o ya es porcentaje)
                                    if pd.notna(valor) and isinstance(valor, (int, float)):
                                        columnas_con_porcentaje.append((col, valor))

                            # La columna correcta es la ÚLTIMA con valores numéricos
                            if columnas_con_porcentaje:
                                ultima_col, porcentaje = columnas_con_porcentaje[-1]  # Tomar la última
                                if isinstance(porcentaje, (int, float)) and porcentaje < 1:
                                    porcentaje = f"{porcentaje * 100:.2f}%"
                                datos["alumnos"][nombre]["porcentaje_asistencia"] = str(porcentaje)
                                st.write(f"   {nombre}: % Asistencia (de columna '{ultima_col}') = {porcentaje}")
                                porcentaje_encontrado = True

                            if not porcentaje_encontrado:
                                st.write(f"   {nombre}: No se encontró columna de % Asistencia válida")

                            for col_excel, codigo_modulo in mapeo_modulos.items():
                                if pd.notna(row.get(col_excel)):
                                    valor = row[col_excel]

                                    if codigo_modulo == 'FCOO03':
                                        datos["alumnos"][nombre]["fcoo03"] = str(valor)
                                        st.write(f"   {nombre}: FCOO03 = {valor}")
                                    elif 'MF' in codigo_modulo:
                                        evaluacion_encontrada = False
                                        if datos_evaluacion and "alumnos" in datos_evaluacion:
                                            nombre_upper = nombre.upper()
                                            for nombre_eval, modulos_eval in datos_evaluacion["alumnos"].items():
                                                if nombre_upper in nombre_eval or nombre_eval in nombre_upper:
                                                    if codigo_modulo in modulos_eval:
                                                        datos["alumnos"][nombre]["modulos_mf"][codigo_modulo] = modulos_eval[codigo_modulo]
                                                        st.write(f"   {nombre}: {codigo_modulo} = {modulos_eval[codigo_modulo]} (Evaluación)")
                                                        evaluacion_encontrada = True
                                                        break

                                        if not evaluacion_encontrada:
                                            datos["alumnos"][nombre]["modulos_mf"][codigo_modulo] = str(valor)
                                            st.write(f"   {nombre}: {codigo_modulo} = {valor} (ASISTENCIA)")

        st.success(f" Datos del Excel leídos: {len(datos['alumnos'])} alumnos encontrados")
        return datos

    except Exception as e:
        st.error(f"Error al leer Excel: {str(e)}")
        import traceback
        st.code(traceback.format_exc())
        return datos

# Funciones de extracción de documentos escaneados
def extraer_datos_certificado_asistencia(texto):
    """Extrae nombres y DNIs de la Hoja de Firmas"""
    datos = {
        "alumnos": [],
        "fecha_inicio": "",
        "curso": ""
    }

    match_curso = re.search(r'Especialidad[:\s]+([^\n\(]+)', texto, re.IGNORECASE)
    if match_curso:
        datos["curso"] = match_curso.group(1).strip()

    fechas = re.findall(r'(\d{2}/\d{2}/\d{4})', texto)
    if fechas:
        datos["fecha_inicio"] = fechas[0]

    patron = r'(\d{8}[A-Z])\s+([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ\s]+?)(?=\s*\d{8}[A-Z]|\n\n|LUNES|MARTES|\Z)'
    matches = re.findall(patron, texto, re.DOTALL)

    palabras_excluir = ['INTERPROS', 'GENERATION', 'OPERACIONES', 'AUXILIARES',
                        'SERVICIOS', 'ADMINISTRATIVOS', 'PRINCIPADO', 'ASTURIAS',
                        'PUBLICO', 'EMPLEO', 'ENTRADA', 'SALIDA', 'MIERCOLES']

    for dni, nombre in matches:
        nombre_limpio = ' '.join(nombre.strip().split())
        es_valido = True
        for palabra in palabras_excluir:
            if palabra in nombre_limpio.upper():
                es_valido = False
                break
        if es_valido and len(nombre_limpio) > 10:
            datos["alumnos"].append((nombre_limpio, dni))

    return datos

def extraer_evaluacion_profesores(texto):
    """Extrae nota final y calificación (Superado/No superado) del documento Evaluación de Profesores"""
    # Esta función ahora es un placeholder - la lectura real se hace en extraer_evaluacion_excel
    datos = {
        "alumnos": {}
    }

    lineas = texto.split('\n')

    nombre_actual = None
    for i, linea in enumerate(lineas):
        linea = linea.strip()

        # Detectar nombres de alumnos (patrón: APELLIDOS, NOMBRE)
        if ',' in linea and len(linea) > 10:
            partes = linea.split(',')
            if len(partes) == 2 and all(p.strip().replace(' ', '').isalpha() for p in partes):
                nombre_actual = linea.upper().strip()
                if nombre_actual not in datos["alumnos"]:
                    datos["alumnos"][nombre_actual] = {}
                    st.write(f"  Alumno detectado en evaluación: {nombre_actual}")

        if nombre_actual:
            linea_upper = linea.upper()

            # Buscar módulo (MF0969_1, etc.)
            if 'MF' in linea_upper and '_' in linea_upper:
                match_modulo = re.search(r'(MF\d{4}_\d)', linea_upper)
                if match_modulo:
                    modulo = match_modulo.group(1)

                    # Buscar nota numérica en las líneas cercanas (buscar número como 10, 9, 8.5, etc.)
                    nota = None
                    calificacion = None

                    # Buscar en las siguientes 5 líneas
                    for j in range(i, min(len(lineas), i+6)):
                        linea_busqueda = lineas[j].strip()

                        # Buscar nota numérica (10, 9, 8.5, etc.)
                        match_nota = re.search(r'\b(\d{1,2}(?:[.,]\d{1,2})?)\b', linea_busqueda)
                        if match_nota and not nota:
                            nota_str = match_nota.group(1).replace(',', '.')
                            try:
                                nota_float = float(nota_str)
                                if 0 <= nota_float <= 10:  # Validar que sea una nota válida
                                    nota = nota_str
                            except:
                                pass

                        # Buscar calificación (Superado/No superado)
                        linea_upper_busqueda = linea_busqueda.upper()
                        calificaciones = ['SUPERADO', 'NO SUPERADO', 'CONVALIDADO', 'CONVALIDA', 'APTO', 'NO APTO', 'EXENTO']
                        for calif in calificaciones:
                            if calif in linea_upper_busqueda and not calificacion:
                                calificacion = calif
                                break

                    # Combinar nota + calificación
                    if nota and calificacion:
                        valor_final = f"{nota} {calificacion}"
                        datos["alumnos"][nombre_actual][modulo] = valor_final
                        st.write(f"   Evaluación: {nombre_actual} - {modulo} → {valor_final}")
                    elif nota:
                        datos["alumnos"][nombre_actual][modulo] = nota
                        st.write(f"   Evaluación: {nombre_actual} - {modulo} → {nota} (sin calificación)")
                    elif calificacion:
                        datos["alumnos"][nombre_actual][modulo] = calificacion
                        st.write(f"   Evaluación: {nombre_actual} - {modulo} → {calificacion} (sin nota)")

    return datos

def extraer_evaluacion_excel(file, verbose=True):
    """Lee el Excel de Evaluación buscando columnas NOTA FINAL y Superado dentro de cada módulo"""
    datos = {
        "alumnos": {}
    }

    try:
        file.seek(0)
        if verbose: st.write(" Leyendo Excel de Evaluación...")

        # Leer TODO el Excel sin encabezados
        df_raw = pd.read_excel(file, sheet_name=0, header=None)

        if verbose: st.write(f"Dimensiones: {df_raw.shape[0]} filas x {df_raw.shape[1]} columnas")

        # 1. Buscar fila con módulos
        fila_modulos = None
        for idx, row in df_raw.iterrows():
            row_text = ' '.join([str(x) for x in row.values if pd.notna(x)])
            if 'MF0969' in row_text or 'MF0970' in row_text:
                fila_modulos = idx
                if verbose: st.write(f" Fila módulos: {fila_modulos}")
                break

        if not fila_modulos:
            return datos

        # 2. Buscar columna de nombres
        col_nombres = None
        for col_idx in range(min(10, df_raw.shape[1])):
            for fila_idx in range(fila_modulos + 1, min(fila_modulos + 15, df_raw.shape[0])):
                valor = df_raw.iloc[fila_idx, col_idx]
                if pd.notna(valor) and ',' in str(valor) and len(str(valor)) > 10:
                    col_nombres = col_idx
                    if verbose: st.write(f" Columna nombres: {col_nombres}")
                    break
            if col_nombres is not None:
                break

        if col_nombres is None:
            return datos

        # 3. Identificar módulos y buscar columnas NOTA FINAL y Superado
        fila_mod = df_raw.iloc[fila_modulos]
        modulos_info = []

        for col_idx, valor in enumerate(fila_mod):
            if pd.notna(valor) and 'MF' in str(valor).upper():
                match = re.search(r'(MF\d{4}_\d)', str(valor).upper())
                if match:
                    modulo = match.group(1)
                    modulos_info.append({"modulo": modulo, "col_inicio": col_idx})

        # Para cada módulo, buscar columnas con NOTA FINAL y Superado
        for i, info in enumerate(modulos_info):
            modulo = info["modulo"]
            col_inicio = info["col_inicio"]
            col_fin = modulos_info[i + 1]["col_inicio"] if i + 1 < len(modulos_info) else df_raw.shape[1]

            nota_col = None
            calif_col = None

            # Buscar en encabezados (filas después de módulos)
            for col_idx in range(col_inicio, col_fin):
                for fila_enc in range(fila_modulos, min(fila_modulos + 5, df_raw.shape[0])):
                    celda = str(df_raw.iloc[fila_enc, col_idx])

                    if 'NOTA' in celda.upper() and 'FINAL' in celda.upper():
                        nota_col = col_idx
                    if 'SUPERADO' in celda.upper():
                        calif_col = col_idx

            if nota_col and calif_col:
                info["nota_col"] = nota_col
                info["calif_col"] = calif_col
                if verbose: st.write(f" {modulo}: nota=col{nota_col}, calif=col{calif_col}")

        # 4. Buscar fila inicio datos
        fila_inicio = None
        for fila_idx in range(fila_modulos + 1, min(fila_modulos + 15, df_raw.shape[0])):
            valor = df_raw.iloc[fila_idx, col_nombres]
            if pd.notna(valor) and ',' in str(valor) and len(str(valor)) > 10:
                fila_inicio = fila_idx
                if verbose: st.write(f" Datos desde fila: {fila_inicio}")
                break

        if not fila_inicio:
            return datos

        # 5. Procesar alumnos
        for fila_idx in range(fila_inicio, df_raw.shape[0]):
            nombre_valor = df_raw.iloc[fila_idx, col_nombres]

            if not pd.notna(nombre_valor):
                continue

            nombre = str(nombre_valor).strip().upper()
            if len(nombre) < 5 or ',' not in nombre:
                continue

            if nombre not in datos["alumnos"]:
                datos["alumnos"][nombre] = {}

            for info in modulos_info:
                if "nota_col" not in info or "calif_col" not in info:
                    continue

                modulo = info["modulo"]
                nota_val = df_raw.iloc[fila_idx, info["nota_col"]]
                calif_val = df_raw.iloc[fila_idx, info["calif_col"]]

                nota = str(nota_val).strip() if pd.notna(nota_val) and str(nota_val) != 'nan' else ""
                calif = str(calif_val).strip() if pd.notna(calif_val) and str(calif_val) != 'nan' else ""

                # Limpiar nota - convertir a entero SIN REDONDEAR (truncar decimales)
                if nota:
                    # Ignorar la palabra "baja" en sí misma, pero procesar el resto
                    if nota.lower() == 'baja':
                        nota = ""
                    else:
                        try:
                            n = float(nota)
                            nota = str(int(n))  # int() trunca, no redondea (9.8 → 9, no 10)
                        except:
                            pass

                # Limpiar calificación también (por si tiene decimales)
                if calif:
                    # Ignorar la palabra "baja" en sí misma
                    if calif.lower() == 'baja':
                        calif = ""
                    else:
                        try:
                            c = float(calif)
                            calif = str(int(c))  # Limpiar decimales de calificación también
                        except:
                            pass

                # Combinar nota y calificación
                if nota and calif:
                    # Si son iguales, mostrar solo UNA vez
                    if nota.lower() == calif.lower():
                        datos["alumnos"][nombre][modulo] = nota
                        if verbose: st.write(f"   {nombre} - {modulo}: {nota} (valores iguales)")
                    else:
                        # Si son diferentes, mostrar ambos
                        datos["alumnos"][nombre][modulo] = f"{nota} {calif}"
                        if verbose: st.write(f"   {nombre} - {modulo}: {nota} {calif}")
                elif nota or calif:
                    # Solo hay uno, mostrarlo
                    valor = nota if nota else calif
                    datos["alumnos"][nombre][modulo] = valor
                    if verbose: st.write(f"   {nombre} - {modulo}: {valor}")

        st.success(f" Procesados {len(datos['alumnos'])} alumnos")
        return datos

    except Exception as e:
        st.error(f"Error: {str(e)}")
        import traceback
        st.code(traceback.format_exc())
        return datos


def llenar_excel_resumen(excel_file, datos_excel, datos_documentos, datos_ctrl=None):
    """Combina datos del Excel, documentos escaneados y CTRL, y rellena la pestaña RESUMEN"""
    try:
        excel_file.seek(0)
        wb = openpyxl.load_workbook(excel_file)

        if "RESUMEN" not in wb.sheetnames:
            st.error("No se encontró la pestaña 'RESUMEN'")
            return None

        ws = wb["RESUMEN"]

        # Leer encabezados
        encabezados = {}
        for col in range(1, ws.max_column + 1):
            valor = ws.cell(row=1, column=col).value
            if valor:
                encabezados[str(valor).strip().lower()] = col

        st.write("**Encabezados detectados en RESUMEN:**")
        st.write(list(encabezados.keys()))

        alumnos_excel = datos_excel.get("alumnos", {})

        if not alumnos_excel:
            st.warning("No se encontraron alumnos en las pestañas del Excel")
            return None

        st.info(f"Se procesarán {len(alumnos_excel)} alumnos")

        # Procesar TODOS los alumnos (sin límite)
        alumnos_lista = list(alumnos_excel.items())
        celdas_escritas = 0

        for i, (nombre, datos_alumno) in enumerate(alumnos_lista):
            fila = 2 + i
            st.write(f"Procesando fila {fila}: {nombre}")

            # ID
            if "id" in encabezados:
                ws.cell(row=fila, column=encabezados["id"], value=i + 1)
                celdas_escritas += 1

            # Nombre completo
            if "nombre completo" in encabezados:
                ws.cell(row=fila, column=encabezados["nombre completo"], value=nombre)
                celdas_escritas += 1

            # DNI - Priorizar del CTRL
            dni = ""
            if datos_ctrl:
                nombre_upper = nombre.upper()
                for nombre_ctrl, datos_ctrl_alumno in datos_ctrl.items():
                    if nombre_upper in nombre_ctrl or nombre_ctrl in nombre_upper:
                        dni_ctrl = datos_ctrl_alumno.get("dni", "")
                        if dni_ctrl:
                            dni = dni_ctrl
                            st.write(f"   DNI del CTRL: {dni}")
                            break

            if not dni:
                dni = datos_alumno.get("dni", "")
                if dni:
                    st.write(f"  ℹ DNI del Excel principal: {dni}")

            if "dni" in encabezados and dni:
                ws.cell(row=fila, column=encabezados["dni"], value=dni)
                celdas_escritas += 1

            # DATOS DEL CTRL
            if datos_ctrl:
                nombre_upper = nombre.upper()
                datos_alumno_ctrl = None

                st.write(f"   DEBUG - Buscando '{nombre_upper}' en CTRL...")
                st.write(f"   DEBUG - Total alumnos en CTRL: {len(datos_ctrl)}")

                for nombre_ctrl, datos_ctrl_alumno in datos_ctrl.items():
                    if nombre_upper in nombre_ctrl or nombre_ctrl in nombre_upper:
                        datos_alumno_ctrl = datos_ctrl_alumno
                        st.write(f"   Encontrado en CTRL: '{nombre_ctrl}'")
                        st.write(f"   DEBUG - Datos del alumno en CTRL:")
                        for key, value in datos_ctrl_alumno.items():
                            st.write(f"      {key}: '{value}'")
                        break

                if not datos_alumno_ctrl:
                    st.write(f"   NO encontrado en CTRL")
                    st.write(f"   DEBUG - Primeros 5 nombres en CTRL:")
                    for i, ctrl_name in enumerate(list(datos_ctrl.keys())[:5]):
                        st.write(f"      {i+1}. '{ctrl_name}'")

                if datos_alumno_ctrl:
                    # Corporación
                    corporacion = datos_alumno_ctrl.get("corporacion_a_clase", "")
                    for enc_key, col in encabezados.items():
                        if "corporacion" in enc_key.lower() or "corporación" in enc_key.lower():
                            if corporacion:
                                ws.cell(row=fila, column=col, value=corporacion)
                                celdas_escritas += 1
                                st.write(f"   Corporación → {corporacion}")
                            break

                    # Baja
                    baja = datos_alumno_ctrl.get("baja", "")
                    for enc_key, col in encabezados.items():
                        if enc_key.lower() == "baja":
                            if baja:
                                ws.cell(row=fila, column=col, value=baja)
                                celdas_escritas += 1
                                st.write(f"   Baja → {baja}")
                            break

                    # Motivo (sin paréntesis)
                    motivo = datos_alumno_ctrl.get("motivo_sin_parentesis", "")
                    for enc_key, col in encabezados.items():
                        if "motivo" in enc_key.lower() and "baja" not in enc_key.lower():
                            if motivo:
                                ws.cell(row=fila, column=col, value=motivo)
                                celdas_escritas += 1
                                st.write(f"   Motivo → {motivo}")
                            break

                    # Baja - Motivo (combinado) con DEBUG
                    baja_fecha_raw = datos_alumno_ctrl.get("baja", "")
                    motivo_sin_parentesis = datos_alumno_ctrl.get("motivo_sin_parentesis", "")

                    # Convertir fecha de baja a formato DD/MM/YYYY si es necesario
                    baja_fecha = ""
                    if baja_fecha_raw:
                        # Si es datetime o date, convertir a string DD/MM/YYYY
                        if isinstance(baja_fecha_raw, (datetime, date)):
                            baja_fecha = baja_fecha_raw.strftime('%d/%m/%Y')
                        else:
                            # Si es string, intentar parsearlo y reformatearlo
                            baja_str = str(baja_fecha_raw)
                            # Si tiene timestamp (00:00:00), quitarlo
                            if '00:00:00' in baja_str:
                                try:
                                    fecha_obj = pd.to_datetime(baja_str)
                                    baja_fecha = fecha_obj.strftime('%d/%m/%Y')
                                except:
                                    baja_fecha = baja_str
                            else:
                                baja_fecha = baja_str

                    st.write(f"   DEBUG Baja-Motivo:")
                    st.write(f"    - baja_fecha_raw: '{baja_fecha_raw}' (tipo: {type(baja_fecha_raw)})")
                    st.write(f"    - baja_fecha (procesada): '{baja_fecha}'")
                    st.write(f"    - motivo_sin_parentesis: '{motivo_sin_parentesis}'")

                    for enc_key, col in encabezados.items():
                        if ("baja" in enc_key.lower() and "motivo" in enc_key.lower()) or \
                           (enc_key.lower() == "baja - motivo"):
                            st.write(f"    - Columna encontrada: '{enc_key}'")
                            # Combinar baja + motivo SIN guion, solo espacio
                            if baja_fecha and motivo_sin_parentesis:
                                baja_motivo_combinado = f"{baja_fecha} {motivo_sin_parentesis}"
                                st.write(f"    - Ambos existen, combinado: '{baja_motivo_combinado}'")
                            elif baja_fecha:
                                baja_motivo_combinado = baja_fecha
                                st.write(f"    - Solo baja: '{baja_motivo_combinado}'")
                            elif motivo_sin_parentesis:
                                baja_motivo_combinado = motivo_sin_parentesis
                                st.write(f"    - Solo motivo: '{baja_motivo_combinado}'")
                            else:
                                baja_motivo_combinado = ""
                                st.write(f"    - Ninguno existe, vacío")

                            if baja_motivo_combinado:
                                ws.cell(row=fila, column=col, value=baja_motivo_combinado)
                                celdas_escritas += 1
                                st.write(f"   Baja - Motivo → '{baja_motivo_combinado}'")
                            break

                    # % Baja Ocupación
                    baja_ocupacion = datos_alumno_ctrl.get("baja_ocupacion", "")
                    for enc_key, col in encabezados.items():
                        if ('baja' in enc_key.lower() and 'ocupacion' in enc_key.lower()) or \
                           ('baja' in enc_key.lower() and 'ocupación' in enc_key.lower()) or \
                           ('%' in enc_key and ('ocupacion' in enc_key.lower() or 'ocupación' in enc_key.lower())):
                            if baja_ocupacion:
                                ws.cell(row=fila, column=col, value=baja_ocupacion)
                                celdas_escritas += 1
                                st.write(f"   % Baja Ocupación → {baja_ocupacion}")
                            break

                    # FECHA DE INCORPORACIÓN - DEBUG COMPLETO
                    fecha_incorporacion = datos_alumno_ctrl.get("fecha_incorporacion", "")

                    st.write(f"   DEBUG FECHA - Valor recibido: {fecha_incorporacion}")
                    st.write(f"   DEBUG FECHA - Tipo: {type(fecha_incorporacion)}")

                    for enc_key, col in encabezados.items():
                        if ('incorporacion' in enc_key.lower() or 'incorporación' in enc_key.lower()) and \
                           ('clase' in enc_key.lower() or 'fecha' in enc_key.lower() or 'sintrafor' in enc_key.lower()):
                            if fecha_incorporacion:
                                celda = ws.cell(row=fila, column=col)

                                st.write(f"   DEBUG - Columna encontrada: '{enc_key}' (col {col})")
                                st.write(f"   DEBUG - Formato actual de la celda ANTES: {celda.number_format}")

                                # Convertir SIEMPRE a STRING en formato DD/MM/YYYY
                                if isinstance(fecha_incorporacion, date):
                                    # Convertir a string DD/MM/YYYY
                                    fecha_string = fecha_incorporacion.strftime('%d/%m/%Y')
                                    celda.value = fecha_string
                                    # Aplicar formato de TEXTO (no fecha)
                                    celda.number_format = '@'  # @ = formato texto
                                    st.write(f"   Fecha Incorporación → '{fecha_string}' (TEXTO)")
                                    st.write(f"   DEBUG - Formato aplicado: '@' (texto)")
                                elif isinstance(fecha_incorporacion, datetime):
                                    # Convertir a string DD/MM/YYYY
                                    fecha_string = fecha_incorporacion.strftime('%d/%m/%Y')
                                    celda.value = fecha_string
                                    celda.number_format = '@'  # @ = formato texto
                                    st.write(f"   Fecha Incorporación → '{fecha_string}' (TEXTO desde datetime)")
                                    st.write(f"   DEBUG - Formato aplicado: '@' (texto)")
                                else:
                                    celda.value = str(fecha_incorporacion)
                                    celda.number_format = '@'
                                    st.write(f"   Fecha Incorporación → '{fecha_incorporacion}' (STRING directo)")
                                    st.write(f"   DEBUG - Formato aplicado: '@' (texto)")

                                st.write(f"   DEBUG - Valor escrito en celda: {celda.value}")
                                st.write(f"   DEBUG - Tipo del valor: {type(celda.value)}")
                                st.write(f"   DEBUG - Formato de la celda DESPUÉS: {celda.number_format}")

                                celdas_escritas += 1
                            break
                else:
                    st.write(f"   No se encontraron datos en CTRL para {nombre}")

            # % ASISTENCIA
            porcentaje = datos_alumno.get("porcentaje_asistencia", "")
            st.write(f"   DEBUG - % Asistencia del alumno: '{porcentaje}'")

            columna_encontrada = False
            for enc_key, col in encabezados.items():
                # Buscar columnas que contengan "asistencia" y "%"
                if "asistencia" in enc_key.lower() and "%" in enc_key:
                    if porcentaje:
                        ws.cell(row=fila, column=col, value=porcentaje)
                        celdas_escritas += 1
                        st.write(f"   % Asistencia → {porcentaje} (columna: '{enc_key}')")
                        columna_encontrada = True
                    break

            if not columna_encontrada and porcentaje:
                st.write(f"   No se encontró columna para % Asistencia")
                st.write(f"   Columnas disponibles con 'asistencia': {[k for k in encabezados.keys() if 'asistencia' in k.lower()]}")


            # ACREDITACION - DEJAR EN BLANCO para rellenar manualmente
            st.write(f"  ℹ ACREDITACION se deja en blanco (para rellenar manualmente)")

            # LIQUIDACION TEORIA - DEJAR EN BLANCO para rellenar manualmente
            st.write(f"  ℹ LIQUIDACION TEORIA se deja en blanco (para rellenar manualmente)")

            # LIQUIDACION EMPRESA - DEJAR EN BLANCO para rellenar manualmente
            st.write(f"  ℹ LIQUIDACION EMPRESA se deja en blanco (para rellenar manualmente)")

            # FCOO03
            fcoo03 = datos_alumno.get("fcoo03", "")
            st.write(f"  DEBUG: fcoo03 = '{fcoo03}'")

            fcoo03_col = None
            for enc_key, col in encabezados.items():
                if 'fcoo' in enc_key.lower() and '03' in enc_key.lower():
                    fcoo03_col = col
                    st.write(f"  DEBUG: Encontrada columna FCOO03: '{enc_key}' (col {col})")
                    break

            if fcoo03_col and fcoo03:
                ws.cell(row=fila, column=fcoo03_col, value=fcoo03)
                celdas_escritas += 1
                st.write(f"   FCOO03 → {fcoo03}")
            elif not fcoo03_col:
                st.write(f"   No se encontró columna FCOO03 en encabezados")
            elif not fcoo03:
                st.write(f"   FCOO03 está vacío para {nombre}")

            # PRL - DEJAR EN BLANCO
            st.write(f"  ℹ PRL se deja en blanco")

            # Módulos MF
            modulos_mf = datos_alumno.get("modulos_mf", {})
            st.write(f"  DEBUG: modulos_mf = {modulos_mf}")

            if modulos_mf:
                for modulo, calificacion in modulos_mf.items():
                    modulo_limpio = str(modulo).strip().upper()
                    st.write(f"  DEBUG: Procesando módulo '{modulo_limpio}' con valor '{calificacion}'")

                    encontrado = False
                    for enc_key, col in encabezados.items():
                        enc_key_upper = enc_key.strip().upper()

                        if modulo_limpio == enc_key_upper or \
                           modulo_limpio.replace('_', '') == enc_key_upper.replace('_', '') or \
                           modulo_limpio in enc_key_upper or \
                           enc_key_upper in modulo_limpio:
                            ws.cell(row=fila, column=col, value=calificacion)
                            celdas_escritas += 1
                            st.write(f"   {modulo} → {calificacion} (columna {col}, encabezado '{enc_key}')")
                            encontrado = True
                            break

                    if not encontrado:
                        st.write(f"   No se encontró columna para módulo '{modulo_limpio}'")
            else:
                st.write(f"  ℹ No hay módulos MF para {nombre}")

            # F.E. - DEJAR EN BLANCO
            st.write(f"  ℹ F.E. se deja en blanco")

        st.success(f" Total: {celdas_escritas} celdas escritas en {len(alumnos_lista)} filas")

        # Guardar
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        return output

    except Exception as e:
        st.error(f"Error: {str(e)}")
        import traceback
        st.code(traceback.format_exc())
        return None

# Estilos CSS
st.markdown("""
<style>
.header {
    width: 100%;
    background: linear-gradient(90deg, #00BCD4 0%, #009688 100%);
    padding: 1rem 2rem;
    display: flex;
    align-items: center;
    justify-content: center;
    border-radius: 0 0 15px 15px;
    box-shadow: 0 4px 15px rgba(0,0,0,0.15);
}
.header img {
    width: 100%;
    max-width: 600px;
    object-fit: contain;
}
.sidebar .sidebar-content {
    background: #f9fafc;
}
.sidebar-title {
    font-size: 1.3rem;
    font-weight: 700;
    color: #009688;
    margin-bottom: 1rem;
}
.main-container {
    background: white;
    padding: 2rem;
    border-radius: 15px;
    box-shadow: 0 4px 15px rgba(0,0,0,0.05);
}
.info-box {
    background: #e0f7fa;
    padding: 1rem;
    border-radius: 10px;
    border-left: 4px solid #00BCD4;
    margin: 1rem 0;
}
</style>
""", unsafe_allow_html=True)

# Encabezado
try:
    logo_b64 = base64.b64encode(open("logo.png", "rb").read()).decode()
    logo_html = f'<img src="data:image/png;base64,{logo_b64}" alt="Smartmind logo">'
except:
    logo_html = '<h1 style="color:white;">Smartmind</h1>'

st.markdown(f'<div class="header">{logo_html}</div>', unsafe_allow_html=True)

# Secciones
secciones = {
    "Captación": "Carga los documentos de captación de alumnos.",
    "Formación Empresa Inicio": "Documentación de inicio de formación en empresa.",
    "Formación Empresa Fin": "Documentación de finalización de formación en empresa.",
    "Evaluación": "Documentos de evaluación del curso.",
    "Cierre Mes": "Documentación de cierre mensual.",
}

if "seccion_actual" not in st.session_state:
    st.session_state.seccion_actual = list(secciones.keys())[0]

# Menú lateral
st.sidebar.markdown('<p class="sidebar-title">Menú</p>', unsafe_allow_html=True)
for nombre in secciones.keys():
    if st.sidebar.button(nombre, key=nombre, use_container_width=True):
        st.session_state.seccion_actual = nombre

st.sidebar.markdown("---")
st.sidebar.info("Usa el menú para navegar entre las secciones del sistema de automatización.")

# Contenido principal
st.markdown(f'<div class="main-container">', unsafe_allow_html=True)

# Header profesional
st.markdown(f"""
<div class="header-container">
    <h1 class="header-title">{st.session_state.seccion_actual}</h1>
    <p class="header-subtitle">{secciones[st.session_state.seccion_actual]}</p>
</div>
""", unsafe_allow_html=True)

if st.session_state.seccion_actual == "Formación Empresa Fin":

    st.markdown("### Fase 1: Cargar Archivos Excel")
    st.markdown('<div class="custom-card"><p>Sube los archivos Excel necesarios para completar el proceso de documentación.</p></div>', unsafe_allow_html=True)

    col_excel1, col_excel2 = st.columns(2)

    with col_excel1:
        st.markdown("**Excel Principal** _(RESUMEN, CALIFICACIONES, ASISTENCIA)_")
        excel_justificacion = st.file_uploader(
            "Cargar Excel con las 3 pestañas",
            type=["xlsx", "xls"],
            key="excel_justificacion"
        )

        if excel_justificacion:
            st.success(" Excel principal cargado")
            try:
                xls = pd.ExcelFile(excel_justificacion)
                st.write(f" Pestañas: {', '.join(xls.sheet_names)}")
            except Exception as e:
                st.warning(f" Error: {str(e)}")

    with col_excel2:
        st.markdown("**Excel CTRL de Alumnos** _(Pestaña CTRL)_")
        excel_ctrl = st.file_uploader(
            "Cargar Excel CTRL de Alumnos",
            type=["xlsx", "xls"],
            key="excel_ctrl",
            help="Este Excel debe contener la pestaña CTRL con información de corporación, baja y motivo"
        )

        if excel_ctrl:
            st.success(" Excel CTRL cargado")
            try:
                xls_ctrl = pd.ExcelFile(excel_ctrl)
                st.write(f" Pestañas: {', '.join(xls_ctrl.sheet_names)}")
                if "CTRL" in xls_ctrl.sheet_names:
                    st.info(" Pestaña CTRL encontrada")
                else:
                    st.warning(" No se encontró la pestaña CTRL")
            except Exception as e:
                st.warning(f" Error: {str(e)}")

    st.markdown("---")
    st.markdown("### Fase 2: Documentos Escaneados (Requeridos)")
    st.markdown('<div class="info-box"> Carga los 4 documentos. Estos son NECESARIOS para extraer datos que NO están en las pestañas del Excel.</div>', unsafe_allow_html=True)

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("**Plan de Evaluación**")
        plan_evaluacion = st.file_uploader(
            "Cargar Plan de Evaluación",
            type=["pdf", "png", "jpg", "jpeg", "bmp", "tiff", "gif", "xlsx", "xls", "docx", "doc", "txt"],
            key="plan_evaluacion"
        )

        st.markdown("**Cronograma**")
        cronograma = st.file_uploader(
            "Cargar Cronograma",
            type=["pdf", "png", "jpg", "jpeg", "bmp", "tiff", "gif", "xlsx", "xls", "docx", "doc", "txt"],
            key="cronograma"
        )

    with col2:
        st.markdown("**Certificado de Asistencia (Hoja de Firmas)**")
        certificado = st.file_uploader(
            "Cargar Hoja de Firmas",
            type=["pdf", "png", "jpg", "jpeg", "bmp", "tiff", "gif", "xlsx", "xls", "docx", "doc", "txt"],
            key="certificado"
        )

        st.markdown("**Evaluación de Profesores**")
        evacuacion = st.file_uploader(
            "Cargar Evaluación",
            type=["pdf", "png", "jpg", "jpeg", "bmp", "tiff", "gif", "xlsx", "xls", "docx", "doc", "txt"],
            key="evacuacion"
        )

    st.markdown("---")
    st.markdown("### Fase 3: Procesamiento Automático")

    documentos_cargados = [plan_evaluacion, cronograma, certificado, evacuacion]
    archivos_totales = sum(1 for doc in documentos_cargados if doc is not None)

    if archivos_totales > 0:
        st.markdown(f'<div class="custom-card"><p><strong>{archivos_totales} de 4</strong> documentos cargados correctamente</p></div>', unsafe_allow_html=True)

    excel_ctrl_cargado = excel_ctrl is not None

    if excel_justificacion and all(documentos_cargados):

        if not excel_ctrl_cargado:
            st.warning("AVISO: Excel CTRL no cargado. Los campos de Corporación, Baja y Motivo quedarán vacíos.")

        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("Procesar y Completar Resumen", type="primary", use_container_width=True):

            # Crear barra de progreso
            progress_bar = st.progress(0)
            status_text = st.empty()

            # Procesar Excel CTRL
            datos_ctrl = None
            if excel_ctrl_cargado:
                status_text.text("Procesando Excel CTRL...")
                progress_bar.progress(10)
                datos_ctrl = leer_datos_ctrl(excel_ctrl)

            # Procesar documentos sin mostrar debug
            datos_documentos = {}
            datos_evaluacion = None

            # Plan de Evaluación
            status_text.text(" Procesando Plan de Evaluación...")
            progress_bar.progress(20)
            if plan_evaluacion:
                texto_plan = procesar_documento(plan_evaluacion)

            # Cronograma
            status_text.text(" Procesando Cronograma...")
            progress_bar.progress(30)
            if cronograma:
                texto_cronograma = procesar_documento(cronograma)

            # Hoja de Firmas
            status_text.text(" Procesando Hoja de Firmas...")
            progress_bar.progress(40)
            if certificado:
                texto_certificado = procesar_documento(certificado)
                if texto_certificado:
                    datos_documentos["certificado_asistencia"] = extraer_datos_certificado_asistencia(texto_certificado)

            # Evaluación de Profesores
            status_text.text("‍ Procesando Evaluación de Profesores...")
            progress_bar.progress(50)
            if evacuacion:
                if evacuacion.type in ["application/vnd.ms-excel", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"]:
                    evacuacion.seek(0)
                    # Silenciar el debug temporalmente
                    import sys
                    from io import StringIO
                    old_stdout = sys.stdout
                    sys.stdout = StringIO()

                    datos_evaluacion = extraer_evaluacion_excel(evacuacion, verbose=False)

                    sys.stdout = old_stdout
                else:
                    texto_evacuacion = procesar_documento(evacuacion)
                    if texto_evacuacion:
                        datos_evaluacion = extraer_evaluacion_profesores(texto_evacuacion)

            # Leer Excel
            status_text.text(" Leyendo datos del Excel principal...")
            progress_bar.progress(70)
            datos_excel = leer_datos_excel(excel_justificacion, datos_evaluacion)

            # Llenar RESUMEN
            status_text.text("Completando resumen...")
            progress_bar.progress(90)

            try:
                excel_justificacion.seek(0)
                excel_bytes = excel_justificacion.read()
                excel_buffer = io.BytesIO(excel_bytes)

                excel_completado = llenar_excel_resumen(excel_buffer, datos_excel, datos_documentos, datos_ctrl)
            except Exception as e:
                st.error(f"ERROR: {str(e)}")
                excel_completado = None

            progress_bar.progress(100)
            status_text.empty()
            progress_bar.empty()

            # Mostrar resultado
            if excel_completado:
                st.markdown("""
                <div class="result-container">
                    <h2 class="result-title">Proceso Completado</h2>
                    <p style="color: #065f46; font-size: 1rem;">El resumen ha sido generado exitosamente con todos los datos integrados.</p>
                </div>
                """, unsafe_allow_html=True)

                # Resumen de procesamiento
                col1, col2, col3, col4 = st.columns(4)

                with col1:
                    st.metric("Alumnos", len(datos_excel.get('alumnos', {})))

                with col2:
                    evaluaciones = sum(len(mods) for mods in datos_evaluacion.get("alumnos", {}).values()) if datos_evaluacion else 0
                    st.metric("Calificaciones", evaluaciones)

                with col3:
                    ctrl_status = "SÍ" if datos_ctrl else "NO"
                    st.metric("Excel CTRL", ctrl_status)

                with col4:
                    docs_procesados = sum([
                        1 if plan_evaluacion else 0,
                        1 if cronograma else 0,
                        1 if certificado else 0,
                        1 if evacuacion else 0
                    ])
                    st.metric("Documentos", f"{docs_procesados}/4")

                st.markdown("<br>", unsafe_allow_html=True)

                # Botones de acción
                col_btn1, col_btn2 = st.columns([2, 1])

                with col_btn1:
                    st.download_button(
                        label="Descargar Excel Completado",
                        data=excel_completado,
                        file_name=f"resumen_completado_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True,
                        type="primary"
                    )

                with col_btn2:
                    if st.button("Vista Previa", use_container_width=True):
                        st.info("Descarga el archivo Excel para ver el contenido completo.")

                # Mensaje informativo
                st.markdown("""
                <div class="custom-card">
                    <h3>Datos Integrados</h3>
                    <p><strong>Información incorporada en el resumen:</strong></p>
                    <ul style="color: var(--text-secondary); line-height: 1.8;">
                        <li>Datos de alumnos del Excel principal y CTRL</li>
                        <li>Fechas de incorporación y motivos de baja</li>
                        <li>DNI y corporación de cada alumno</li>
                        <li>Porcentajes de asistencia (MF+FCOO+FE)</li>
                        <li>Calificaciones de módulos (MF) de la evaluación</li>
                        <li>Calificaciones de FCOO03</li>
                    </ul>
                    <p style="margin-top: 1rem; font-weight: 500;">Descarga el archivo y revisa que toda la información sea correcta.</p>
                </div>
                """, unsafe_allow_html=True)

            else:
                st.error("ERROR: Hubo un error al completar el Excel. Por favor, revisa los archivos cargados.")

    elif excel_justificacion and archivos_totales > 0:
        st.warning("AVISO: Por favor, carga TODOS los documentos necesarios para continuar")
        st.info(f"Documentos faltantes: {4 - archivos_totales}")

else:
    st.info("Esta sección está en desarrollo. Por el momento, solo 'Formación Empresa Fin' tiene funcionalidad de procesamiento automático.")

st.markdown("</div>", unsafe_allow_html=True)